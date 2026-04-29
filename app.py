import csv
import io
import json
import os
import uuid

from dotenv import load_dotenv
load_dotenv(dotenv_path=".claude/.env")

from flask import Flask, jsonify, render_template, request, send_file, session
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt
import openpyxl

from extractor import PROMPT_VARIANTS, extract_entities
from eval_runner import run_evals, load_logs

app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET", "dev-secret-key")

# Server-side storage keyed by session ID — avoids 4KB cookie size limit
_doc_store: dict[str, str] = {}
_extraction_store: dict[str, dict] = {}

MODELS = [
    # OpenAI direct
    {"value": "openai::gpt-4.1-nano",         "label": "OpenAI — gpt-4.1-nano (fastest, cheapest)"},
    {"value": "openai::gpt-4.1-mini",         "label": "OpenAI — gpt-4.1-mini"},
    {"value": "openai::gpt-4.1",              "label": "OpenAI — gpt-4.1"},
    {"value": "openai::gpt-4o-mini",          "label": "OpenAI — gpt-4o-mini"},
    {"value": "openai::gpt-4o",               "label": "OpenAI — gpt-4o"},
    {"value": "openai::gpt-5",                "label": "OpenAI — gpt-5"},
    {"value": "openai::o1-mini",              "label": "OpenAI — o1-mini (reasoning)"},
    {"value": "openai::o3-mini",              "label": "OpenAI — o3-mini (reasoning)"},
    # OpenRouter — open source / free tier
    {"value": "openrouter::mistralai/mistral-7b-instruct",      "label": "OpenRouter — Mistral 7B (free)"},
    {"value": "openrouter::meta-llama/llama-3-8b-instruct",     "label": "OpenRouter — Llama 3 8B (free)"},
    {"value": "openrouter::google/gemma-3-12b-it:free",         "label": "OpenRouter — Gemma 3 12B (free)"},
    # OpenRouter — frontier models
    {"value": "openrouter::anthropic/claude-3-5-haiku",         "label": "OpenRouter — Claude 3.5 Haiku"},
    {"value": "openrouter::anthropic/claude-sonnet-4-5",        "label": "OpenRouter — Claude Sonnet 4.5"},
    {"value": "openrouter::google/gemini-2.0-flash-001",        "label": "OpenRouter — Gemini 2.0 Flash"},
    {"value": "openrouter::openai/gpt-4o-mini",                 "label": "OpenRouter — GPT-4o-mini (via OR)"},
]


def _ensure_session():
    if "sid" not in session:
        session["sid"] = str(uuid.uuid4())


def _extract_text(file_bytes: bytes, filename: str) -> str:
    if filename.lower().endswith(".pdf"):
        reader = PdfReader(io.BytesIO(file_bytes))
        text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
    elif filename.lower().endswith((".txt", ".md")):
        text = file_bytes.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type: {filename}")
    if len(text) > 50_000:
        text = text[:50_000] + "\n[...document truncated at 50,000 chars]"
    return text


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.route("/")
def index():
    _ensure_session()
    return render_template("index.html", models=MODELS, prompt_variants=list(PROMPT_VARIANTS.keys()))


@app.route("/upload", methods=["POST"])
def upload():
    _ensure_session()
    sid = session["sid"]
    if "document" not in request.files:
        return jsonify({"ok": False, "error": "No file provided"}), 400
    f = request.files["document"]
    try:
        text = _extract_text(f.read(), f.filename)
        _doc_store[sid] = text
        _extraction_store.pop(sid, None)
        return jsonify({"ok": True, "chars": len(text)})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 400


@app.route("/extract", methods=["POST"])
def extract():
    _ensure_session()
    sid = session["sid"]
    doc_text = _doc_store.get(sid)
    if not doc_text:
        return jsonify({"error": "No document uploaded"}), 400

    body = request.get_json(force=True)
    raw_model = body.get("model", "openai::gpt-4o-mini")
    provider, model = raw_model.split("::", 1)
    prompt_variant = body.get("prompt_variant", "v2_detailed")

    result = extract_entities(doc_text, provider, model, prompt_variant)
    _extraction_store[sid] = result
    return jsonify(result)


@app.route("/download/<fmt>")
def download(fmt: str):
    sid = session.get("sid")
    extraction = _extraction_store.get(sid) if sid else None
    if not extraction or "error" in extraction:
        return "No extraction available", 400

    if fmt == "csv":
        return _download_csv(extraction)
    elif fmt == "excel":
        return _download_excel(extraction)
    elif fmt == "word":
        return _download_word(extraction)
    return "Unknown format", 400


@app.route("/evals")
def evals_page():
    return render_template("evals.html", models=MODELS, prompt_variants=list(PROMPT_VARIANTS.keys()))


@app.route("/evals/logs")
def evals_logs():
    return jsonify(load_logs())


@app.route("/evals/run", methods=["POST"])
def evals_run():
    body = request.get_json(force=True)
    raw_model = body.get("model", "openai::gpt-4o-mini")
    provider, model = raw_model.split("::", 1)
    prompt_variant = body.get("prompt_variant", "v2_detailed")

    results = run_evals(provider=provider, model=model, prompt_variant=prompt_variant)
    return jsonify(results)


# ---------------------------------------------------------------------------
# Export helpers
# ---------------------------------------------------------------------------

def _download_csv(data: dict):
    buf = io.StringIO()
    w = csv.writer(buf)

    w.writerow(["PARTIES"])
    w.writerow(["Name"])
    for p in data.get("parties", []):
        w.writerow([p])
    w.writerow([])

    w.writerow(["DATES"])
    w.writerow(["Label", "Value"])
    for d in data.get("dates", []):
        w.writerow([d.get("label", ""), d.get("value", "")] if isinstance(d, dict) else ["", d])
    w.writerow([])

    w.writerow(["AMOUNTS"])
    w.writerow(["Label", "Value"])
    for a in data.get("amounts", []):
        w.writerow([a.get("label", ""), a.get("value", "")] if isinstance(a, dict) else ["", a])
    w.writerow([])

    w.writerow(["OBLIGATIONS"])
    w.writerow(["Obligation"])
    for o in data.get("obligations", []):
        w.writerow([o])
    w.writerow([])

    w.writerow(["GOVERNING LAW"])
    w.writerow([data.get("governing_law", "")])

    buf.seek(0)
    return send_file(
        io.BytesIO(buf.getvalue().encode()),
        mimetype="text/csv",
        as_attachment=True,
        download_name="contract_entities.csv",
    )


def _download_excel(data: dict):
    wb = openpyxl.Workbook()

    def add_sheet(title, headers, rows):
        ws = wb.create_sheet(title=title)
        ws.append(headers)
        for row in rows:
            ws.append(row)

    add_sheet("Parties", ["Name"], [[p] for p in data.get("parties", [])])
    add_sheet("Dates", ["Label", "Value"],
              [[d.get("label", ""), d.get("value", "")] if isinstance(d, dict) else ["", d]
               for d in data.get("dates", [])])
    add_sheet("Amounts", ["Label", "Value"],
              [[a.get("label", ""), a.get("value", "")] if isinstance(a, dict) else ["", a]
               for a in data.get("amounts", [])])
    add_sheet("Obligations", ["Obligation"], [[o] for o in data.get("obligations", [])])
    add_sheet("Governing Law", ["Jurisdiction"], [[data.get("governing_law", "")]])

    if "Sheet" in wb.sheetnames:
        del wb["Sheet"]

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                     as_attachment=True, download_name="contract_entities.xlsx")


def _download_word(data: dict):
    doc = Document()
    doc.add_heading("Contract Entity Extraction", 0)

    def add_table(heading, headers, rows):
        doc.add_heading(heading, level=1)
        if not rows:
            doc.add_paragraph("None found.")
            return
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                table.rows[ri + 1].cells[ci].text = str(val)
        doc.add_paragraph("")

    add_table("Parties", ["Name"], [[p] for p in data.get("parties", [])])
    add_table("Dates", ["Label", "Value"],
              [[d.get("label", ""), d.get("value", "")] if isinstance(d, dict) else ["", d]
               for d in data.get("dates", [])])
    add_table("Amounts", ["Label", "Value"],
              [[a.get("label", ""), a.get("value", "")] if isinstance(a, dict) else ["", a]
               for a in data.get("amounts", [])])
    add_table("Obligations", ["Obligation"], [[o] for o in data.get("obligations", [])])

    doc.add_heading("Governing Law", level=1)
    doc.add_paragraph(str(data.get("governing_law") or "Not specified"))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True, download_name="contract_entities.docx")


if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5001)
